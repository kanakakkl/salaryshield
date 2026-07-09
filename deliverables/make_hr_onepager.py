"""Build a one-page HR/CHRO sales sheet for SalaryShield (.docx)."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

VIOLET = RGBColor(0x6D, 0x28, 0xD9)
DARK = RGBColor(0x1F, 0x29, 0x37)
GREY = RGBColor(0x55, 0x5D, 0x6B)
EMER = RGBColor(0x0F, 0x76, 0x5C)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

doc = Document()
sec = doc.sections[0]
sec.top_margin = Inches(0.5); sec.bottom_margin = Inches(0.5)
sec.left_margin = Inches(0.6); sec.right_margin = Inches(0.6)

st = doc.styles["Normal"].font
st.name = "Segoe UI"; st.size = Pt(10); st.color.rgb = DARK


def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd"); sh.set(qn("w:val"), "clear"); sh.set(qn("w:fill"), hexcolor)
    tcPr.append(sh)


def para(text, size=10, color=DARK, bold=False, italic=False, align=None, after=4, before=0):
    p = doc.add_paragraph(); r = p.add_run(text)
    r.font.size = Pt(size); r.font.color.rgb = color; r.font.bold = bold; r.font.italic = italic; r.font.name = "Segoe UI"
    p.paragraph_format.space_after = Pt(after); p.paragraph_format.space_before = Pt(before)
    if align is not None: p.alignment = align
    return p


def bullet(text, bold_lead=None, size=9.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    if bold_lead:
        r = p.add_run(bold_lead); r.font.bold = True; r.font.name = "Segoe UI"; r.font.size = Pt(size); r.font.color.rgb = DARK
    r2 = p.add_run(text); r2.font.name = "Segoe UI"; r2.font.size = Pt(size); r2.font.color.rgb = DARK
    return p

# ---------------- Header band ----------------
tbl = doc.add_table(rows=1, cols=1); tbl.autofit = True
cell = tbl.rows[0].cells[0]
shade(cell, "6D28D9")
cp = cell.paragraphs[0]; cp.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = cp.add_run("SalaryShield for HR & People Teams")
r.font.size = Pt(20); r.font.bold = True; r.font.color.rgb = WHITE; r.font.name = "Segoe UI"
cp2 = cell.add_paragraph()
r2 = cp2.add_run("A comp-planning copilot: turn your annual hike cycle into a location-aware, ROI-justified decision.")
r2.font.size = Pt(11); r2.font.color.rgb = WHITE; r2.font.italic = True; r2.font.name = "Segoe UI"
cell.paragraphs[0].paragraph_format.space_before = Pt(4)
cp2.paragraph_format.space_after = Pt(4)

para("", after=6)

# ---------------- The problem, in HR terms ----------------
para("The problem you're already feeling", 13, VIOLET, bold=True, after=4)
para("Your annual hike is a flat percentage. But 8% in Mumbai (9.1% inflation) is a real pay cut, while 8% in Pune "
     "(6.8%) is a real raise. Nobody budgets for that gap today — so your best people in your most expensive cities "
     "leave first, and each exit costs roughly 1.5x annual salary to replace.", 10, after=8)

# ---------------- Value table ----------------
para("What SalaryShield gives your team", 13, VIOLET, bold=True, after=4)
rows = [
    ("HR priority", "SalaryShield feature", "What it shows you"),
    ("Retention", "Attrition-risk KPI + inflation-lag chart", "Which cities / tenure bands are in the high-alert zone right now"),
    ("Budget planning", "Live compensation budget simulator", "Slide a corrective hike % -> exact cost vs. retention ROI, instantly"),
    ("CFO justification", "ROI multiplier + net savings", "\"This hike prevents this much replacement cost\" — a raise as an investment, not a cost"),
    ("Location strategy", "City-level inflation heatmap", "Which office locations are becoming the most expensive to retain talent in"),
    ("Pay equity / compliance", "Fairness audit (gender gap, tenure equity, remediation budget)", "A built-in, always-current pay-equity audit trail"),
]
t = doc.add_table(rows=len(rows), cols=3); t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.LEFT
widths = [Inches(1.5), Inches(2.3), Inches(3.0)]
for i, r3 in enumerate(rows):
    for j, val in enumerate(r3):
        c = t.rows[i].cells[j]; c.width = widths[j]
        c.paragraphs[0].text = ""
        run = c.paragraphs[0].add_run(val)
        run.font.name = "Segoe UI"; run.font.size = Pt(8.7)
        run.font.bold = (i == 0)
        run.font.color.rgb = WHITE if i == 0 else DARK
        if i == 0:
            shade(c, "6D28D9")
        elif i % 2 == 0:
            shade(c, "F1ECFB")

para("", after=8)

# ---------------- Two-column: why it works / honest caveats ----------------
tw = doc.add_table(rows=1, cols=2)
c1, c2 = tw.rows[0].cells
c1.width = Inches(3.4); c2.width = Inches(3.4)

p = c1.paragraphs[0]; p.text = ""
r = p.add_run("Why it works"); r.font.bold = True; r.font.size = Pt(11); r.font.color.rgb = EMER; r.font.name = "Segoe UI"
for bl, tx in [
    ("Replacement-cost framing: ", "uses ~1.5x annual salary per replaced hire — the number that gets a CFO's attention."),
    ("A live planning tool, not a static report: ", "move the slider in the room, in the meeting, and get the answer."),
    ("Defensible, not generic: ", "\"Mumbai needs more than Pune\" is backed by real cost-of-living data, not a guess."),
    ("Same numbers on both sides: ", "employees see the identical CLII data in their own portal, cutting negotiation friction."),
]:
    pp = c1.add_paragraph(style="List Bullet"); pp.paragraph_format.space_after = Pt(3)
    rr = pp.add_run(bl); rr.font.bold = True; rr.font.size = Pt(9); rr.font.name = "Segoe UI"; rr.font.color.rgb = DARK
    rr2 = pp.add_run(tx); rr2.font.size = Pt(9); rr2.font.name = "Segoe UI"; rr2.font.color.rgb = DARK

p2 = c2.paragraphs[0]; p2.text = ""
r2 = p2.add_run("Honest scope today"); r2.font.bold = True; r2.font.size = Pt(11); r2.font.color.rgb = RGBColor(0xB4, 0x53, 0x09); r2.font.name = "Segoe UI"
for bl, tx in [
    ("Illustrative models: ", "attrition-risk and replacement-cost coefficients are not yet fitted to your own exit data."),
    ("Sample workforce data: ", "pilot deployments would import real headcount via CSV or an HRIS connector."),
    ("No auth/multi-tenant yet: ", "a production rollout needs SSO and role-based access (HR vs. employee view)."),
    ("Live data feeds: ", "MOSPI / RBI / rental indices are modeled today; live ingestion is the next milestone.")
]:
    pp = c2.add_paragraph(style="List Bullet"); pp.paragraph_format.space_after = Pt(3)
    rr = pp.add_run(bl); rr.font.bold = True; rr.font.size = Pt(9); rr.font.name = "Segoe UI"; rr.font.color.rgb = DARK
    rr2 = pp.add_run(tx); rr2.font.size = Pt(9); rr2.font.name = "Segoe UI"; rr2.font.color.rgb = DARK

para("", after=8)

# ---------------- Pilot ask ----------------
tbl2 = doc.add_table(rows=1, cols=1)
cell2 = tbl2.rows[0].cells[0]; shade(cell2, "F1ECFB")
cp3 = cell2.paragraphs[0]
rr = cp3.add_run("Pilot proposal: "); rr.font.bold = True; rr.font.size = Pt(10); rr.font.color.rgb = VIOLET; rr.font.name = "Segoe UI"
rr2 = cp3.add_run("run SalaryShield against one business unit's real comp data for one appraisal cycle. "
                   "Compare the corrective-hike recommendation against your existing flat-percentage plan, "
                   "and measure the delta in projected attrition risk and budget efficiency.")
rr2.font.size = Pt(10); rr2.font.color.rgb = DARK; rr2.font.name = "Segoe UI"
cell2.paragraphs[0].paragraph_format.space_before = Pt(6)
cell2.paragraphs[0].paragraph_format.space_after = Pt(6)

para("", after=6)
para("SalaryShield  |  React 19 + Vite 8  |  Google Gemini-powered Copilot  |  City-Level Inflation Index (CLII)",
     8.5, GREY, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)

doc.save("C:/Users/ADMIN/kkl-workspace/salaryshield/deliverables/SalaryShield_HR_OnePager.docx")
print("saved SalaryShield_HR_OnePager.docx")
