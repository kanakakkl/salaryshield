"""Build the SalaryShield hackathon design document (.docx)."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

VIOLET = RGBColor(0x6D, 0x28, 0xD9)
DARK = RGBColor(0x1F, 0x29, 0x37)
GREY = RGBColor(0x55, 0x5D, 0x6B)
EMER = RGBColor(0x0F, 0x76, 0x5C)

doc = Document()
# Base style
st = doc.styles["Normal"].font
st.name = "Segoe UI"; st.size = Pt(10.5); st.color.rgb = DARK

for name, size in (("Heading 1", 17), ("Heading 2", 13)):
    f = doc.styles[name].font
    f.name = "Segoe UI"; f.size = Pt(size); f.bold = True; f.color.rgb = VIOLET


def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd"); sh.set(qn("w:val"), "clear"); sh.set(qn("w:fill"), hexcolor)
    tcPr.append(sh)


def para(text, size=10.5, color=DARK, bold=False, italic=False, align=None, after=6, before=0):
    p = doc.add_paragraph(); r = p.add_run(text)
    r.font.size = Pt(size); r.font.color.rgb = color; r.font.bold = bold; r.font.italic = italic; r.font.name = "Segoe UI"
    p.paragraph_format.space_after = Pt(after); p.paragraph_format.space_before = Pt(before)
    if align is not None: p.alignment = align
    return p


def bullet(text, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_lead:
        r = p.add_run(bold_lead); r.font.bold = True; r.font.name = "Segoe UI"; r.font.size = Pt(10.5); r.font.color.rgb = DARK
    r2 = p.add_run(text); r2.font.name = "Segoe UI"; r2.font.size = Pt(10.5); r2.font.color.rgb = DARK
    p.paragraph_format.space_after = Pt(4)
    return p

# ---------------- Cover ----------------
para("SalaryShield", 34, VIOLET, bold=True, after=2)
para("AI-Driven, Inflation-Adaptive Compensation Intelligence", 15, GREY, bold=True, after=14)
para("Design & Solution Document", 12, DARK, bold=True, after=2)
para("Hackathon submission — mapped to the 10 scoring parameters", 10.5, GREY, italic=True, after=18)
para("Quantify real-wage erosion. Justify fair corrections. Retain talent. SalaryShield is a two-sided "
     "platform that turns invisible city-level inflation into a fair, data-backed compensation decision for "
     "both employees and employers — powered by a City-Level Inflation Index (CLII) and a GenAI Copilot "
     "grounded in the user's own numbers.", 11, DARK, after=14)

# Tech stack table
para("Technology Stack", 13, VIOLET, bold=True, after=6)
tech = [
    ("Layer", "Technology"),
    ("Frontend", "React 19, Vite 8, JavaScript (JSX)"),
    ("UI / Design", "CSS variables + glassmorphism design system, lucide-react icons"),
    ("State & Persistence", "React Hooks (useState / useEffect), browser localStorage"),
    ("Domain / Model", "src/lib/inflation.js — CLII model, analyzeBudget(), savingsSeries(), analyzeWorkforce(), break-even engine"),
    ("AI (implemented)", "src/lib/llm.js — Google Gemini 2.0 Flash, live API calls grounded on user budget + CLII, multi-turn"),
    ("AI (roadmap)", "Server-side LLM proxy for key security + rate limiting, response streaming"),
    ("Data (roadmap)", "MOSPI CPI, RBI, MagicBricks / 99acres, Zomato / Swiggy ingestion pipeline"),
    ("Tooling", "oxlint, Vite build (1,780 modules, ~86 KB gzip)"),
]
tbl = doc.add_table(rows=len(tech), cols=2); tbl.style = "Table Grid"; tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
tbl.columns[0].width = Inches(2.3); tbl.columns[1].width = Inches(4.2)
for i, (a, b) in enumerate(tech):
    c0, c1 = tbl.rows[i].cells
    c0.width = Inches(2.3); c1.width = Inches(4.2)
    for cell, text, bold in ((c0, a, True), (c1, b, i == 0)):
        cell.paragraphs[0].text = ""
        r = cell.paragraphs[0].add_run(text)
        r.font.name = "Segoe UI"; r.font.size = Pt(9.5); r.font.bold = bold or (i == 0)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF) if i == 0 else DARK
    if i == 0:
        shade(c0, "6D28D9"); shade(c1, "6D28D9")
    else:
        shade(c0, "F1ECFB")

doc.add_page_break()

# ---------------- Architecture ----------------
para("Solution Architecture", 17, VIOLET, bold=True, after=6)
para("SalaryShield is a layered, model-driven single-page application. The presentation layer hosts five "
     "modules; all of them read a single domain model (src/lib/inflation.js), so every metric and every AI "
     "answer is computed from one source of truth. Client state is held in React hooks and persisted to "
     "localStorage, which is how the user's budget flows from the Budget Planner into the Employee Portal, "
     "the Negotiation Coach, and the Copilot in real time. The AI layer is live today — the Copilot calls "
     "Google Gemini directly, grounded in the same model. Only the underlying data-ingestion pipeline "
     "(replacing modeled CLII figures with live MOSPI/RBI/rent feeds) remains on the roadmap.", 10.5, after=10)
doc.add_picture("C:/Users/ADMIN/kkl-workspace/salaryshield/deliverables/architecture.png", width=Inches(6.5))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
para("Figure 1 — SalaryShield layered architecture (LIVE and ROADMAP layers tagged).", 9, GREY, italic=True,
     align=WD_ALIGN_PARAGRAPH.CENTER, after=14)

doc.add_page_break()

# ---------------- 10 parameters ----------------
para("Evaluation Against the 10 Scoring Parameters", 17, VIOLET, bold=True, after=10)

sections = [
    ("1. Problem Definition",
     "Nominal salary is not real income. Inflation across India's tech hubs silently erodes purchasing power, "
     "yet the national CPI under-represents how urban tech workers actually spend.",
     [("", "Tech-hub inflation is high and uneven — Hyderabad 8.4%, Mumbai 9.1%, Bengaluru 7.9%, Pune 6.8%."),
      ("", "National MOSPI CPI under-weights the urban tech-worker basket: IT-corridor rent, education, food delivery, commute."),
      ("", "Employees cannot quantify their real-wage loss or defend a corrective hike with evidence."),
      ("", "Employers lose talent to inflation lag without city-level visibility — each exit costs roughly Rs.37.5L to replace."),
      ("Outcome: ", "mis-priced pay, avoidable attrition, and unfair compensation across geographies.")]),

    ("2. Solution Architecture",
     "A layered SPA with a single domain model shared by every module (see Figure 1).",
     [("Presentation: ", "five React modules — Employer Console, Employee Portal, My Budget & Inflation, CLII Engine, GenAI Copilot."),
      ("Domain model: ", "src/lib/inflation.js is the single source of truth — CLII data, analyzeBudget(), savingsSeries(), break-even and personal-inflation engine."),
      ("State: ", "React hooks + localStorage; the user's budget syncs live across Planner -> Employee Portal -> Coach -> Copilot."),
      ("AI layer: ", "the GenAI Copilot (src/lib/llm.js) is live today, calling Google Gemini directly through the same model boundary.")]),

    ("3. Engineering / Build Quality",
     "The app is clean, verifiable, and production-shaped.",
     [("", "React 19 + Vite 8 SPA; one component per module; CSS-variable design system (glassmorphism)."),
      ("DRY: ", "domain logic extracted into src/lib/inflation.js and reused by all five modules, including the Employer Console's attrition/ROI/fairness math."),
      ("Testable: ", "pure, deterministic functions (analyzeBudget, savingsSeries, analyzeWorkforce) — unit-test-ready."),
      ("Cross-module sync: ", "changing city or salary in the Budget Planner propagates live to the Employee Portal, Coach and Copilot within one second."),
      ("Verified: ", "production build passes (1,780 modules, ~86 KB gzip); oxlint clean; localStorage persistence; live in-browser checks with zero console errors.")]),

    ("4. AI Appropriateness",
     "AI is applied only where it adds unique value.",
     [("", "Hard numbers — inflation, break-even, ROI — stay in deterministic code for accuracy and auditability."),
      ("", "The LLM handles what it is uniquely good at: reasoning over those numbers and expressing them in plain language."),
      ("Principle: ", "the engine computes, the AI explains — no AI-for-AI's-sake and no hallucinated math."),
      ("", "The output is genuinely language-shaped: personalized advice and ready-to-send negotiation scripts.")]),

    ("5. AI Implementation",
     "A real, live LLM integration — grounded, multi-turn, and consistent.",
     [("", "The Copilot calls Google Gemini 2.0 Flash directly (src/lib/llm.js) — a live model call, not a scripted response."),
      ("RAG-style grounding: ", "a system prompt injects the user's live budget, computed metrics, and the full CLII dataset for every city, so answers cite the user's real rupee figures rather than model priors."),
      ("Multi-turn: ", "full conversation history is passed on every call, so follow-up questions stay in context."),
      ("Consistency: ", "Budget Planner, Employee Portal, Negotiation Coach and Copilot all read the same model — one data story."),
      ("Roadmap: ", "a server-side proxy for API key security and rate limiting, plus streaming responses.")]),

    ("6. AI Impact",
     "Invisible loss becomes a clear, immediate action.",
     [("Employee: ", "surfaces a hidden ~Rs.1.03L/yr loss and generates a copy-paste, data-backed pitch — improving raise success."),
      ("Employer: ", "converts attrition risk into ROI live in the Employer Console — retain a key hire and avoid ~Rs.37.5L in replacement cost, recomputed instantly as the correction hike is adjusted."),
      ("Insight: ", "personalized inflation (e.g. 7.7% vs a 8.4% city average) that no generic calculator provides."),
      ("", "Data becomes a decision in seconds, for both the individual and the organization.")]),

    ("7. Business Relevance",
     "A large, recurring, two-sided need.",
     [("", "Addressable users: 5M+ Indian IT professionals and every enterprise running India delivery centers."),
      ("", "The pain is universal and recurring — it resurfaces at every appraisal cycle."),
      ("", "Maps directly to top HR priorities: retention, pay equity, and compensation-budget optimization."),
      ("", "Two-sided design gives both employee and employer a reason to adopt.")]),

    ("8. Benefits & Viability",
     "A clear model with low cost and a real data moat.",
     [("Revenue: ", "freemium employee tier for acquisition; B2B SaaS (per-seat / per-audit) from employers."),
      ("Moat: ", "the proprietary CLII index compounds in value as coverage and history grow."),
      ("Economics: ", "a static SPA plus serverless LLM calls keeps infrastructure lean."),
      ("Path: ", "MVP is built; remaining pieces (live feeds, HRMS integration, auth) are known and incremental.")]),

    ("9. Working Demo",
     "A running, interactive MVP — not slideware.",
     [("Step 1 — Budget Planner: ", "enter salary, city & expenses; see live personal inflation, a 12-month savings-depletion chart, and the break-even raise; push rent up to trigger the deficit alert."),
      ("Step 2 — Negotiation Coach: ", "the pitch auto-cites those exact numbers; one-click copy."),
      ("Step 3 — GenAI Copilot: ", "'Am I underpaid?' triggers a real Google Gemini call that answers with the same live figures."),
      ("Step 4 — Employer Console: ", "move the corrective-hike slider and watch attrition risk, ROI, and the pay-fairness audit recompute instantly.")]),

    ("10. Differentiation",
     "Not another salary calculator.",
     [("", "Personal, city-level, expense-weighted inflation — not a one-size national average."),
      ("", "Two-sided: employee guidance and employer workforce economics in a single product."),
      ("", "AI output is an actionable artifact — a negotiation script — not just a number."),
      ("", "A single, consistent data story flows across every module; the CLII index is defensible IP.")]),
]

for title, lead, pts in sections:
    para(title, 13, VIOLET, bold=True, after=3, before=8)
    para(lead, 10.5, GREY, italic=True, after=4)
    for bl, tx in pts:
        bullet(tx, bold_lead=bl or None)

doc.add_page_break()
para("Roadmap", 17, VIOLET, bold=True, after=6)
for bl, tx in [("Now: ", "MVP live — CLII model, five modules, live Employer Console, budget engine, Gemini-powered Copilot."),
               ("Next: ", "live data ingestion (MOSPI / RBI / rent / consumer basket) and the CLII computation pipeline."),
               ("Next: ", "server-side LLM proxy for API key security and rate limiting, plus streaming responses."),
               ("Then: ", "HRMS integration, authentication, multi-user, and a historical CLII data moat.")]:
    bullet(tx, bold_lead=bl)
para("", after=6)
para("The one-line pitch: SalaryShield turns invisible inflation into a fair, data-backed pay decision — "
     "for the employee and the employer.", 11, VIOLET, bold=True, before=6)

doc.save("C:/Users/ADMIN/kkl-workspace/salaryshield/deliverables/SalaryShield_Design_Document.docx")
print("saved SalaryShield_Design_Document.docx")
